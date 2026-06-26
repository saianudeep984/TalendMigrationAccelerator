"""
app/reports/intelligence_exports.py
Thin export utility used by upgrade_advisor_dashboard, framework_intelligence_dashboard,
and migration_runbook_generator.
"""
from __future__ import annotations

import io
import json
from typing import Any, Dict, Optional


def export_markdown_report(
    data: Dict[str, Any],
    markdown: str,
    title: str,
    fmt: str = "md",
    extra_sheets: Optional[Dict[str, Any]] = None,
) -> bytes | str:
    """
    Export a report in the requested format.

    fmt values
    ----------
    "md"   -> returns markdown string
    "json" -> returns JSON bytes
    "html" -> returns HTML bytes
    "pdf"  -> returns PDF bytes (falls back to HTML if weasyprint absent)
    "docx" -> returns DOCX bytes (falls back to HTML if python-docx absent)
    "xlsx" -> returns Excel bytes (falls back to JSON if openpyxl absent)
    """
    fmt = (fmt or "md").lower()

    if fmt == "md":
        return markdown

    if fmt == "json":
        return json.dumps(data, indent=2, default=str).encode("utf-8")

    if fmt == "html":
        html = _markdown_to_html(title, markdown)
        return html.encode("utf-8")

    if fmt == "pdf":
        try:
            import weasyprint  # type: ignore
            html = _markdown_to_html(title, markdown)
            return weasyprint.HTML(string=html).write_pdf()
        except Exception:
            return _markdown_to_html(title, markdown).encode("utf-8")

    if fmt == "docx":
        try:
            from docx import Document  # type: ignore
            doc = Document()
            doc.add_heading(title, 0)
            for line in markdown.splitlines():
                doc.add_paragraph(line)
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except Exception:
            return _markdown_to_html(title, markdown).encode("utf-8")

    if fmt == "xlsx":
        try:
            import openpyxl  # type: ignore
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = title[:31]
            ws.append(["Report", title])
            ws.append([])
            if extra_sheets:
                for sheet_name, rows in extra_sheets.items():
                    ws2 = wb.create_sheet(title=sheet_name[:31])
                    if rows and isinstance(rows[0], dict):
                        headers = list(rows[0].keys())
                        ws2.append(headers)
                        for row in rows:
                            ws2.append([row.get(h, "") for h in headers])
                    else:
                        for row in rows:
                            ws2.append(row if isinstance(row, list) else [row])
            buf = io.BytesIO()
            wb.save(buf)
            return buf.getvalue()
        except Exception:
            return json.dumps(data, indent=2, default=str).encode("utf-8")

    # fallback
    return markdown


# ── internal helpers ──────────────────────────────────────────────────────────

def _markdown_to_html(title: str, markdown: str) -> str:
    try:
        import markdown as md_lib  # type: ignore
        body = md_lib.markdown(markdown, extensions=["tables", "fenced_code"])
    except Exception:
        # no markdown lib — wrap in <pre>
        body = f"<pre>{markdown}</pre>"

    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>{title}</title>
<style>
  body {{ font-family: Arial, sans-serif; max-width: 960px; margin: 40px auto; padding: 0 20px; color: #222; }}
  h1,h2,h3 {{ color: #1a3a5c; }}
  table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
  th,td {{ border: 1px solid #ccc; padding: 8px 12px; text-align: left; }}
  th {{ background: #f0f4f8; }}
  pre,code {{ background: #f5f5f5; padding: 4px 8px; border-radius: 4px; }}
</style>
</head>
<body>
<h1>{title}</h1>
{body}
</body>
</html>"""
