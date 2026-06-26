import html
import logging
import os
import zipfile
from typing import Any, Dict, List, Mapping

logger = logging.getLogger(__name__)


def markdown_fragment_to_html(text: str) -> str:
    """Canonical markdown -> HTML *fragment* converter (table/list/code aware,
    no <html>/<head> wrapper). Shared by every renderer that embeds markdown
    sections inside a larger document (report packs, multi-section exports).
    Promoted from report_pack_generator._markdown_to_html (F5.5/F5.6) so all
    callers use one implementation."""
    out: List[str] = []
    in_code = False
    in_ul = False
    in_table = False
    table_rows: List[List[str]] = []

    def flush_table() -> None:
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return
        data_rows = [
            r for r in table_rows
            if not all(set(c.replace("-", "").replace(":", "").replace(" ", "")) == set() for c in r)
        ]
        if not data_rows:
            in_table = False
            table_rows = []
            return
        html_rows = []
        for i, row in enumerate(data_rows):
            tag = "th" if i == 0 else "td"
            cells = "".join(f"<{tag}>{html.escape(c)}</{tag}>" for c in row)
            tr_style = ' style="background:#f8fafc"' if i % 2 == 0 and i > 0 else ""
            html_rows.append(f"<tr{tr_style}>{cells}</tr>")
        out.append(
            '<table style="border-collapse:collapse;width:100%;margin:8px 0;font-size:10pt">'
            + "".join(html_rows)
            + "</table>"
        )
        in_table = False
        table_rows = []

    for line in text.splitlines():
        stripped = line.strip()

        if stripped.startswith("|"):
            if in_ul:
                out.append("</ul>")
                in_ul = False
            in_table = True
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            table_rows.append(cells)
            continue

        if in_table:
            flush_table()
            table_rows = []

        if stripped.startswith("```"):
            if in_ul:
                out.append("</ul>")
                in_ul = False
            if in_code:
                out.append("</pre>")
                in_code = False
            else:
                out.append("<pre>")
                in_code = True
            continue
        if in_code:
            out.append(html.escape(line))
            continue
        if stripped.startswith("### "):
            if in_ul:
                out.append("</ul>")
                in_ul = False
            out.append(f"<h3>{html.escape(stripped[4:])}</h3>")
        elif stripped.startswith("## "):
            if in_ul:
                out.append("</ul>")
                in_ul = False
            out.append(f"<h2>{html.escape(stripped[3:])}</h2>")
        elif stripped.startswith("# "):
            if in_ul:
                out.append("</ul>")
                in_ul = False
            out.append(f"<h2>{html.escape(stripped[2:])}</h2>")
        elif stripped.startswith("- "):
            if not in_ul:
                out.append("<ul>")
                in_ul = True
            out.append(f"<li>{html.escape(stripped[2:])}</li>")
        elif stripped == "":
            if in_ul:
                out.append("</ul>")
                in_ul = False
            out.append("")
        else:
            if in_ul:
                out.append("</ul>")
                in_ul = False
            out.append(f"<p>{html.escape(stripped)}</p>")

    if in_table:
        flush_table()
    if in_ul:
        out.append("</ul>")
    if in_code:
        out.append("</pre>")
    return "\n".join(out)


def resolve_source_version(all_jobs, repository_path: str = None) -> str:
    """Canonical source-version resolver, shared by excel_export.py and
    json_export.py (F5.5 — was duplicated verbatim in both)."""
    from app.repository.repository_type_detector import RepositoryTypeDetector
    source_version = "UNKNOWN"
    if repository_path:
        try:
            source_version = RepositoryTypeDetector().extract_source_version_from_path(repository_path)
        except Exception:
            pass
    if source_version in (None, "UNKNOWN"):
        for job_entry in all_jobs:
            job = job_entry.get("job_data", job_entry) if isinstance(job_entry, dict) else job_entry
            sv = job.get("source_version")
            if sv and sv not in (None, "UNKNOWN"):
                source_version = sv
                break
    return source_version or "UNKNOWN"


def markdown_table(rows, headers):
    if not rows:
        return ""
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(str(row.get(header, "")) for header in headers) + " |")
    return "\n".join(lines)


def markdown_to_html(markdown: str, title: str) -> str:
    body = []
    in_code = False
    code_lang = ""
    code_lines = []
    has_mermaid = False

    def flush_code():
        nonlocal has_mermaid
        content = "\n".join(code_lines)
        if code_lang == "mermaid":
            has_mermaid = True
            body.append(f"<pre class='mermaid'>{html.escape(content)}</pre>")
        else:
            body.append(f"<pre><code>{html.escape(content)}</code></pre>")

    for line in markdown.splitlines():
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
                code_lang = ""
                code_lines = []
            else:
                in_code = True
                code_lang = line[3:].strip().lower()
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("# "):
            body.append(f"<h1>{html.escape(line[2:])}</h1>")
        elif line.startswith("## "):
            body.append(f"<h2>{html.escape(line[3:])}</h2>")
        elif line.startswith("### "):
            body.append(f"<h3>{html.escape(line[4:])}</h3>")
        elif line.startswith("- "):
            body.append(f"<li>{html.escape(line[2:])}</li>")
        elif line.startswith("|"):
            body.append(f"<pre>{html.escape(line)}</pre>")
        elif not line.strip():
            body.append("<br/>")
        else:
            body.append(f"<p>{html.escape(line)}</p>")
    if in_code:
        flush_code()
    mermaid_script = (
        "<script type='module'>"
        "import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.esm.min.mjs';"
        "mermaid.initialize({startOnLoad:true,theme:'base',themeVariables:{"
        "primaryColor:'#FFFFFF',primaryBorderColor:'#6366F1',lineColor:'#7DD3FC',"
        "primaryTextColor:'#3F3F46',fontFamily:'Arial'}});"
        "</script>"
        if has_mermaid else ""
    )
    return (
        "<!doctype html><html><head><meta charset='utf-8'/>"
        f"<title>{html.escape(title)}</title>"
        "<style>body{font-family:Arial,sans-serif;margin:40px;line-height:1.45}"
        "pre{background:#f6f8fa;padding:6px;border:1px solid #ddd;overflow:auto}"
        ".mermaid{background:#fff;border:1px solid #e5e7eb;border-radius:6px;padding:16px}</style>"
        f"{mermaid_script}</head><body>" + "\n".join(body) + "</body></html>"
    )


def write_pdf(path: str, text: str) -> str:
    safe = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")[:12000]
    stream = f"BT /F1 9 Tf 40 780 Td ({safe}) Tj ET"
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        f"<< /Length {len(stream)} >>\nstream\n{stream}\nendstream".encode("utf-8"),
    ]
    pdf = bytearray(b"%PDF-1.4\n")
    offsets = []
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{index} 0 obj\n".encode("ascii"))
        pdf.extend(obj)
        pdf.extend(b"\nendobj\n")
    xref = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode("ascii"))
    for offset in offsets:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    pdf.extend(f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode("ascii"))
    with open(path, "wb") as handle:
        handle.write(bytes(pdf))
    return path


def write_executive_summary_pdf(path: str, metrics: Mapping[str, Any]) -> str:
    """Write a chart-capable executive PDF using ReportLab vector graphics."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            KeepTogether,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
        from reportlab.graphics.charts.barcharts import HorizontalBarChart, VerticalBarChart
        from reportlab.graphics.charts.piecharts import Pie
        from reportlab.graphics.shapes import Drawing, String
    except Exception:
        logger.exception("Failed to import ReportLab; falling back to simple PDF writer.")
        return write_pdf(path, metrics.get("markdown", "Executive Summary"))

    palette = {
        "navy": colors.HexColor("#0F172A"),
        "muted": colors.HexColor("#64748B"),
        "blue": colors.HexColor("#2563EB"),
        "purple": colors.HexColor("#7C3AED"),
        "teal": colors.HexColor("#14B8A6"),
        "green": colors.HexColor("#22C55E"),
        "amber": colors.HexColor("#F59E0B"),
        "red": colors.HexColor("#EF4444"),
        "surface": colors.HexColor("#F8FAFC"),
        "border": colors.HexColor("#CBD5E1"),
    }

    doc = SimpleDocTemplate(
        path,
        pagesize=A4,
        rightMargin=36,
        leftMargin=36,
        topMargin=34,
        bottomMargin=34,
        title="Executive Summary",
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="TmaTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=24,
        leading=28,
        textColor=palette["navy"],
        spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        name="TmaSection",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        textColor=palette["navy"],
        spaceBefore=14,
        spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        name="TmaBody",
        parent=styles["BodyText"],
        fontSize=9,
        leading=13,
        textColor=palette["muted"],
    ))
    styles.add(ParagraphStyle(
        name="TmaCardLabel",
        parent=styles["BodyText"],
        fontSize=7,
        leading=9,
        textColor=palette["muted"],
        alignment=1,
    ))
    styles.add(ParagraphStyle(
        name="TmaCardValue",
        parent=styles["BodyText"],
        fontName="Helvetica-Bold",
        fontSize=15,
        leading=18,
        textColor=palette["navy"],
        alignment=1,
    ))

    portfolio = metrics.get("portfolio", {})
    scores = metrics.get("scores", {})
    complexity = metrics.get("complexity_breakdown", {})
    cloud = metrics.get("cloud_breakdown", {})
    risks = metrics.get("risk_breakdown", {})

    story = [
        Paragraph("Executive Summary", styles["TmaTitle"]),
        Paragraph(
            "Talend estate migration readiness, cloud suitability, risk exposure, automation coverage, and delivery effort.",
            styles["TmaBody"],
        ),
        Spacer(1, 14),
    ]

    cards = [
        ("Total Jobs", portfolio.get("total_jobs", 0)),
        ("Readiness", f"{scores.get('migration_readiness_score', 0)}%"),
        ("Cloud Score", f"{scores.get('cloud_readiness_score', 0)}%"),
        ("Complexity", f"{scores.get('repository_complexity_score', 0)}%"),
        ("High/Critical Risks", portfolio.get("high_risk", 0)),
        ("Estimated Effort", f"{portfolio.get('estimated_weeks', 'N/A')} wks"),
    ]
    card_cells = [
        [
            Paragraph(str(label), styles["TmaCardLabel"]),
            Paragraph(str(value), styles["TmaCardValue"]),
        ]
        for label, value in cards
    ]
    kpi_table = Table(
        [[cell[0] for cell in card_cells[:3]], [cell[1] for cell in card_cells[:3]]],
        colWidths=[1.65 * inch] * 3,
        rowHeights=[0.23 * inch, 0.36 * inch],
    )
    kpi_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), palette["surface"]),
        ("BOX", (0, 0), (-1, -1), 0.5, palette["border"]),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.white),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    kpi_table_2 = Table(
        [[cell[0] for cell in card_cells[3:]], [cell[1] for cell in card_cells[3:]]],
        colWidths=[1.65 * inch] * 3,
        rowHeights=[0.23 * inch, 0.36 * inch],
    )
    kpi_table_2.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), palette["surface"]),
        ("BOX", (0, 0), (-1, -1), 0.5, palette["border"]),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.white),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.extend([kpi_table, Spacer(1, 6), kpi_table_2])

    score_labels = [
        "Migration readiness",
        "Cloud readiness",
        "Repository complexity",
        "Documentation readiness",
        "Testing readiness",
    ]
    score_values = [
        scores.get("migration_readiness_score", 0),
        scores.get("cloud_readiness_score", 0),
        scores.get("repository_complexity_score", 0),
        scores.get("documentation_readiness_score", 0),
        scores.get("testing_readiness_score", 0),
    ]
    score_drawing = Drawing(500, 185)
    score_chart = HorizontalBarChart()
    score_chart.x = 150
    score_chart.y = 24
    score_chart.height = 130
    score_chart.width = 310
    score_chart.data = [score_values]
    score_chart.valueAxis.valueMin = 0
    score_chart.valueAxis.valueMax = 100
    score_chart.valueAxis.valueStep = 20
    score_chart.categoryAxis.categoryNames = score_labels
    score_chart.bars[0].fillColor = palette["blue"]
    score_chart.barSpacing = 6
    score_drawing.add(String(0, 165, "Readiness Scorecard", fontName="Helvetica-Bold", fontSize=12, fillColor=palette["navy"]))
    score_drawing.add(score_chart)
    story.extend([Spacer(1, 14), score_drawing])

    chart_row = []
    if complexity:
        drawing = Drawing(235, 190)
        chart = Pie()
        chart.x = 52
        chart.y = 28
        chart.width = 128
        chart.height = 128
        chart.data = list(complexity.values())
        chart.labels = [str(k) for k in complexity.keys()]
        for idx, color in enumerate([palette["green"], palette["amber"], palette["red"], palette["purple"], palette["teal"]]):
            if idx < len(chart.slices):
                chart.slices[idx].fillColor = color
        drawing.add(String(0, 170, "Jobs by Complexity", fontName="Helvetica-Bold", fontSize=11, fillColor=palette["navy"]))
        drawing.add(chart)
        chart_row.append(drawing)
    if cloud:
        drawing = Drawing(235, 190)
        chart = VerticalBarChart()
        chart.x = 34
        chart.y = 36
        chart.height = 110
        chart.width = 170
        chart.data = [list(cloud.values())]
        chart.categoryAxis.categoryNames = [str(k) for k in cloud.keys()]
        chart.valueAxis.valueMin = 0
        chart.bars[0].fillColor = palette["teal"]
        drawing.add(String(0, 170, "Cloud Readiness Distribution", fontName="Helvetica-Bold", fontSize=11, fillColor=palette["navy"]))
        drawing.add(chart)
        chart_row.append(drawing)
    if chart_row:
        story.extend([Paragraph("Portfolio Charts", styles["TmaSection"]), Table([chart_row], colWidths=[250] * len(chart_row))])

    risk_rows = [["Risk Signal", "Count"]]
    for label, value in risks.items():
        risk_rows.append([str(label), str(value)])
    if len(risk_rows) == 1:
        risk_rows.append(["No high or critical risk signals", "0"])
    risk_table = Table(risk_rows, colWidths=[360, 90])
    risk_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), palette["navy"]),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("BACKGROUND", (0, 1), (-1, -1), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, palette["border"]),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, palette["surface"]]),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.extend([
        Paragraph("Risk & Decision Narrative", styles["TmaSection"]),
        Paragraph(
            metrics.get(
                "narrative",
                "Approve remediation for high-impact findings, validate import readiness, and proceed through a controlled Talend 8 migration path.",
            ),
            styles["TmaBody"],
        ),
        Spacer(1, 8),
        KeepTogether(risk_table),
    ])

    doc.build(story)
    return path


def write_docx(path: str, title: str, markdown: str) -> str:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title_par = doc.add_heading(title, level=0)
        title_par.alignment = WD_ALIGN_PARAGRAPH.LEFT

        in_table = False
        table_rows: list = []

        def _flush_table():
            nonlocal in_table, table_rows
            if not table_rows:
                in_table = False
                return
            # Determine column count from first data row (skip separator)
            data_rows = [r for r in table_rows if not set(r.replace("|", "").replace("-", "").replace(" ", "")) == set()]
            if not data_rows:
                in_table = False
                table_rows = []
                return
            cols = [c.strip() for c in data_rows[0].strip("|").split("|")]
            ncols = len(cols)
            tbl = doc.add_table(rows=1, cols=ncols)
            tbl.style = "Table Grid"
            hdr_cells = tbl.rows[0].cells
            for i, col in enumerate(cols):
                hdr_cells[i].text = col
                for run in hdr_cells[i].paragraphs[0].runs:
                    run.bold = True
            for row_str in data_rows[1:]:
                cells = [c.strip() for c in row_str.strip("|").split("|")]
                row_cells = tbl.add_row().cells
                for i, cell_val in enumerate(cells[:ncols]):
                    row_cells[i].text = cell_val
            doc.add_paragraph()
            in_table = False
            table_rows = []

        for line in markdown.splitlines():
            # Table accumulation
            if line.startswith("|"):
                in_table = True
                table_rows.append(line)
                continue
            if in_table:
                _flush_table()

            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("```"):
                pass  # skip fence markers
            elif line.strip():
                doc.add_paragraph(line)
            else:
                doc.add_paragraph()

        if in_table:
            _flush_table()

        doc.save(path)

    except ImportError:
        # Fallback: minimal raw OOXML (python-docx not installed)
        lines = [title] + markdown.splitlines()
        paragraphs = "".join(
            "<w:p><w:r><w:t xml:space=\"preserve\">"
            + html.escape(ln)
            + "</w:t></w:r></w:p>"
            for ln in lines
        )
        document = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f"<w:body>{paragraphs}</w:body></w:document>"
        )
        content_types = (
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/word/document.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            "</Types>"
        )
        rels = (
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
            'Target="word/document.xml"/></Relationships>'
        )
        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as docx:
            docx.writestr("[Content_Types].xml", content_types)
            docx.writestr("_rels/.rels", rels)
            docx.writestr("word/document.xml", document)

    return path


def export_document(output_dir: str, basename: str, title: str, markdown: str) -> Dict[str, str]:
    os.makedirs(output_dir, exist_ok=True)
    paths = {
        "markdown": os.path.join(output_dir, f"{basename}.md"),
        "html": os.path.join(output_dir, f"{basename}.html"),
        "pdf": os.path.join(output_dir, f"{basename}.pdf"),
        "docx": os.path.join(output_dir, f"{basename}.docx"),
    }
    with open(paths["markdown"], "w", encoding="utf-8") as handle:
        handle.write(markdown)
    with open(paths["html"], "w", encoding="utf-8") as handle:
        handle.write(markdown_to_html(markdown, title))
    write_pdf(paths["pdf"], markdown)
    write_docx(paths["docx"], title, markdown)
    return paths
