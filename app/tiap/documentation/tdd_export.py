"""
TMA TDD Export Engine
Assembles the full Technical Design Document (all sections) for a job and exports
to Markdown, HTML, DOCX, and PDF — filenames follow <Job_Name>_TDD.<ext>.
Includes Graphviz-generated diagrams embedded as SVG (HTML) and PNG (DOCX/PDF).
"""
from __future__ import annotations
import base64
import io
import os
import re
from typing import Any

from app.tiap.documentation.export_utils import markdown_to_html, write_docx
from app.tiap.documentation import tdd_sections
from app.tiap.testing.testing_architecture import build_testing_architecture
from app.tiap.migration_assessment.migration_assessment import build_migration_assessment
from app.tiap.exec_summary.exec_summary import build_executive_summary
from app.parser.source_target_extractor import build_source_target_inventory


def _safe_job_name(job_name: str) -> str:
    return re.sub(r"[^A-Za-z0-9_]+", "_", job_name or "Unknown").strip("_") or "Unknown"


def _bullets(items: list[str]) -> str:
    return "\n".join(f"- {i}" for i in items) if items else "- None"


# ── Diagram Builders ──────────────────────────────────────────────────────────

_PIPELINE_STAGES = ["Source", "Validation", "Transformation", "Enrichment", "Error Handling", "Target"]
_STAGE_COLORS    = {
    "Source": "#0ea5e9", "Validation": "#f59e0b", "Transformation": "#6366f1",
    "Enrichment": "#10b981", "Error Handling": "#dc2626", "Target": "#64748b",
}
_STAGE_KEYWORDS  = {
    "Source":         {"tFileInputDelimited","tFileInputExcel","tFileInputJSON","tFileInputXML",
                       "tMysqlInput","tOracleInput","tS3Get","tFTPGet","tFixedFlowInput",
                       "tDBInput","tPostgresqlInput","tMSSqlInput","tSalesforceInput","tKafkaInput"},
    "Validation":     {"tFilterRow","tSchemaComplianceCheck","tUniqRow","tCheckEmpty","tAssert"},
    "Transformation": {"tMap","tJava","tJavaRow","tJavaFlex","tAggregateRow","tSortRow",
                       "tNormalize","tDenormalize","tReplaceList","tConvertType"},
    "Enrichment":     {"tJoin","tLookupInput","tSalesforceGetDeleted","tHTTPClient"},
    "Error Handling": {"tDie","tWarn","tLogCatcher","tStatCatcher","tFlowMeter","tDie"},
    "Target":         {"tFileOutputDelimited","tFileOutputExcel","tFileOutputJSON","tFileOutputXML",
                       "tMysqlOutput","tOracleOutput","tS3Put","tFTPPut","tDBOutput",
                       "tPostgresqlOutput","tMSSqlOutput","tSalesforceOutput","tKafkaOutput",
                       "tLogRow","tMysqlClose","tMysqlConnection"},
}

def _classify_stage(ctype: str) -> str:
    for stage, ctypes in _STAGE_KEYWORDS.items():
        if ctype in ctypes:
            return stage
    return "Transformation"


def _dot_to_svg(dot_src: str) -> str:
    """Legacy stub — not called when pure-Python renderer is active."""
    return '<p style="color:#64748b">Diagram unavailable (Graphviz not installed).</p>'


def _dot_to_png_b64(dot_src: str) -> str | None:
    """Legacy stub — not called when pure-Python renderer is active."""
    return None


def _build_pipeline_dot(job_data: dict) -> str:
    """High-level pipeline diagram: Source→Validation→Transformation→…→Target."""
    components = job_data.get("components", [])
    stage_counts = {s: 0 for s in _PIPELINE_STAGES}
    for c in components:
        stage_counts[_classify_stage(c.get("component_type", ""))] += 1

    dot = [
        'digraph G {',
        '  rankdir=LR;',
        '  node [shape=box, style=filled, fontsize=12, fontname="Arial", '
        '        fontcolor="white", color="none", width=1.8, height=0.6];',
        '  edge [color="#94a3b8", penwidth=2];',
    ]
    for s in _PIPELINE_STAGES:
        label = f"{s}\\n({stage_counts[s]} components)"
        dot.append(f'  "{s}" [label="{label}", fillcolor="{_STAGE_COLORS[s]}"];')
    for a, b in zip(_PIPELINE_STAGES[:-1], _PIPELINE_STAGES[1:]):
        dot.append(f'  "{a}" -> "{b}";')
    dot.append("}")
    return "\n".join(dot)


def _build_flow_dot(job_data: dict) -> str:
    """Detailed job flow diagram: components grouped by subjob with connections."""
    components  = job_data.get("components", [])
    connections = job_data.get("connections", [])
    if not components:
        return ""

    TRIGGER_RE = re.compile(r"SUBJOB|RUN_IF|COMPONENT_OK|COMPONENT_ERROR|ITERATE|^OK$|^ERROR$", re.I)

    # Union-Find grouping by data-flow links
    parent = {c.get("unique_name") or c.get("component_type", f"c{i}"): None
               for i, c in enumerate(components)}
    def find(x):
        while parent[x]: x = parent[x]
        return x
    def union(a, b):
        ra, rb = find(a), find(b)
        if ra != rb: parent[ra] = rb

    data_edges, trigger_edges = [], []
    for conn in connections:
        s = conn.get("source", ""); t = conn.get("target", ""); ct = conn.get("connector_type", "")
        if s in parent and t in parent:
            if TRIGGER_RE.search(ct or ""):
                trigger_edges.append((s, t, ct))
            else:
                data_edges.append((s, t, ct))
                union(s, t)

    # Build subjobs
    groups: dict[str, list[str]] = {}
    for name in parent:
        root = find(name)
        groups.setdefault(root, []).append(name)
    subjob_list = list(groups.values())

    type_lookup = {(c.get("unique_name") or c.get("component_type")): c.get("component_type", "")
                   for c in components}

    dot = [
        'digraph G {',
        '  rankdir=TB; fontname="Arial";',
        '  node [shape=box, style=filled, fontsize=9, fontname="Arial", '
        '        fillcolor="#6366f1", fontcolor="white", color="none"];',
        '  edge [fontsize=7];',
    ]
    for idx, members in enumerate(subjob_list):
        dot.append(f'  subgraph cluster_{idx} {{')
        dot.append(f'    label="Subjob {idx+1}"; style=rounded; color="#94a3b8"; fontsize=10; fontname="Arial";')
        for m in members:
            label = f"{m}\\n({type_lookup.get(m, '')})"
            dot.append(f'    "{m}" [label="{label}"];')
        dot.append("  }")
    for s, t, ct in data_edges:
        dot.append(f'  "{s}" -> "{t}" [color="#0ea5e9", label="{ct}", fontcolor="#0ea5e9"];')
    for s, t, ct in trigger_edges:
        dot.append(f'  "{s}" -> "{t}" [color="#dc2626", style=dashed, label="{ct}", fontcolor="#dc2626"];')
    dot.append("}")
    return "\n".join(dot)


def _build_dependency_dot(job_data: dict, all_jobs_data: list[dict] | None = None) -> str:
    """Job orchestration dependency diagram."""
    job_name = job_data.get("job_name", "Unknown")
    # Try to get child/parent info from dependencies
    # all_jobs_data is optional list of job_data dicts for cross-job edges
    child_jobs = [c.get("unique_name", "") for c in job_data.get("components", [])
                  if c.get("component_type") == "tRunJob"]

    dot = [
        'digraph G {',
        '  rankdir=LR; fontname="Arial";',
        '  node [shape=box, style=filled, fontsize=11, fontname="Arial", '
        '        fontcolor="white", color="none"];',
    ]
    dot.append(f'  "{job_name}" [fillcolor="#6366f1"];')
    for child in child_jobs:
        if child:
            dot.append(f'  "{child}" [fillcolor="#94a3b8"];')
            dot.append(f'  "{job_name}" -> "{child}";')
    if not child_jobs:
        dot.append(f'  "No child jobs" [fillcolor="#e2e8f0", fontcolor="#64748b"];')
        dot.append(f'  "{job_name}" -> "No child jobs" [style=dashed];')
    dot.append("}")
    return "\n".join(dot)


# ── HTML with embedded SVGs ───────────────────────────────────────────────────

_HTML_STYLE = """
<style>
  body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.5; color: #0f172a; }
  h1 { color: #0f172a; border-bottom: 3px solid #6366f1; padding-bottom: 8px; }
  h2 { color: #1e293b; margin-top: 32px; border-left: 4px solid #6366f1;
       padding-left: 12px; }
  h3 { color: #334155; margin-top: 18px; }
  table { border-collapse: collapse; width: 100%; margin: 12px 0; font-size: 13px; }
  th { background: #0f172a; color: #fff; padding: 7px 10px; text-align: left; }
  td { border: 1px solid #cbd5e1; padding: 6px 10px; }
  tr:nth-child(even) { background: #f8fafc; }
  li { margin: 4px 0; }
  ul { padding-left: 20px; }
  .diagram-box { border: 1px solid #e2e8f0; border-radius: 8px; padding: 16px;
                 margin: 16px 0; background: #fff; overflow-x: auto; }
  .diagram-label { font-size: 12px; color: #64748b; font-weight: 600;
                   text-transform: uppercase; letter-spacing: .06em; margin-bottom: 8px; }
  .kpi-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(140px,1fr));
              gap: 12px; margin: 16px 0; }
  .kpi-card { background: #f1f5f9; border-radius: 8px; padding: 14px; text-align: center; }
  .kpi-val  { font-size: 28px; font-weight: 700; color: #6366f1; }
  .kpi-lbl  { font-size: 11px; color: #64748b; margin-top: 4px; }
  pre { background: #f6f8fa; padding: 8px; border: 1px solid #ddd;
        border-radius: 4px; overflow: auto; font-size: 12px; }
  .rag-green  { background:#16a34a; color:#fff; border-radius:12px; padding:2px 10px; font-size:12px; }
  .rag-amber  { background:#d97706; color:#fff; border-radius:12px; padding:2px 10px; font-size:12px; }
  .rag-red    { background:#dc2626; color:#fff; border-radius:12px; padding:2px 10px; font-size:12px; }
</style>
"""

def _md_bold(text: str) -> str:
    """Convert **bold** markers to <strong> in a plain text snippet."""
    return re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)

def _section_html(title: str, content_html: str) -> str:
    return f'<h2>{title}</h2>\n{content_html}\n'

def _ul(items: list[str]) -> str:
    if not items:
        return "<ul><li>None</li></ul>"
    return "<ul>" + "".join(f"<li>{_md_bold(i)}</li>" for i in items) + "</ul>"

def _table_html(headers: list[str], rows: list[list[str]]) -> str:
    th = "".join(f"<th>{h}</th>" for h in headers)
    body = "".join(
        "<tr>" + "".join(f"<td>{c}</td>" for c in row) + "</tr>"
        for row in rows
    )
    return f"<table><thead><tr>{th}</tr></thead><tbody>{body}</tbody></table>"

def _kpi_grid(*pairs) -> str:
    cards = "".join(
        f'<div class="kpi-card"><div class="kpi-val">{val}</div>'
        f'<div class="kpi-lbl">{lbl}</div></div>'
        for lbl, val in pairs
    )
    return f'<div class="kpi-grid">{cards}</div>'

def _diagram_box(label: str, svg: str) -> str:
    return (f'<div class="diagram-box">'
            f'<div class="diagram-label">{label}</div>'
            f'{svg}</div>')

def _rag_badge(rag: str) -> str:
    cls = {"GREEN": "rag-green", "AMBER": "rag-amber", "RED": "rag-red"}.get(rag, "rag-amber")
    return f'<span class="{cls}">{rag}</span>'


def build_tdd_sections_html(job_data: dict) -> list[tuple[str, str]]:
    """Build the rich HTML TDD as a list of (## heading, section_html) pairs —
    same content/diagrams/KPI-grids as build_tdd_html, but kept addressable
    per-section so callers (e.g. the section picker in tdd_page.py) can filter
    to a subset of sections without losing styling or diagrams. The heading
    strings match the `## Heading` markers used by build_tdd_markdown / the
    section picker's _SECTION_HEADINGS table, so the two stay in lockstep."""
    job_name   = job_data.get("job_name", "Unknown")
    components = job_data.get("components", [])
    inv        = build_source_target_inventory(job_data)

    sections     = tdd_sections.generate_all_sections(job_data)
    testing      = build_testing_architecture(job_data)
    assessment   = build_migration_assessment(job_data)
    exec_summary = build_executive_summary(job_data)

    # Diagrams — pure-Python renderer (no Graphviz / native binary required)
    import re as _re
    from app.tiap.documentation import diagram_renderer as _dr

    # Pipeline diagram
    _stage_counts = {s: 0 for s in _PIPELINE_STAGES}
    for _c in job_data.get("components", []):
        _stage_counts[_classify_stage(_c.get("component_type", ""))] += 1
    pipeline_svg = _dr.build_pipeline_svg(job_data, _PIPELINE_STAGES, _stage_counts)

    # Flow diagram
    _TRIGGER_RE2 = _re.compile(r"SUBJOB|RUN_IF|COMPONENT_OK|COMPONENT_ERROR|ITERATE|^OK$|^ERROR$", _re.I)
    _components2  = job_data.get("components", [])
    _connections2 = job_data.get("connections", [])
    _parent2 = {(_c.get("unique_name") or _c.get("component_type", f"c{_i}")): None
               for _i, _c in enumerate(_components2)}
    def _find2(x):
        while _parent2[x]: x = _parent2[x]
        return x
    def _union2(a, b):
        ra, rb = _find2(a), _find2(b)
        if ra != rb: _parent2[ra] = rb
    _data_edges2, _trigger_edges2 = [], []
    for _conn2 in _connections2:
        _s2 = _conn2.get("source", ""); _t2 = _conn2.get("target", ""); _ct2 = _conn2.get("connector_type", "")
        if _s2 in _parent2 and _t2 in _parent2:
            if _TRIGGER_RE2.search(_ct2 or ""):
                _trigger_edges2.append((_s2, _t2, _ct2))
            else:
                _data_edges2.append((_s2, _t2, _ct2))
                _union2(_s2, _t2)
    _groups2: dict = {}
    for _name2 in _parent2:
        _root2 = _find2(_name2)
        _groups2.setdefault(_root2, []).append(_name2)
    _subjob_list2 = list(_groups2.values())
    _type_lookup2 = {(_c.get("unique_name") or _c.get("component_type")): _c.get("component_type", "")
                    for _c in _components2}
    flow_svg = _dr.build_flow_svg(job_data, _subjob_list2, _type_lookup2, _data_edges2, _trigger_edges2)

    # Dependency diagram
    _job_name2 = job_data.get("job_name", "Unknown")
    _child_jobs2 = [_c.get("unique_name", "") for _c in job_data.get("components", [])
                   if _c.get("component_type") == "tRunJob"]
    dep_svg = _dr.build_dependency_svg(_job_name2, _child_jobs2)

    out: list[tuple[str, str]] = []

    # ── Executive Summary ────────────────────────────────────────────────────
    parts: list[str] = ["<h3>Business Summary</h3>"]
    for line in exec_summary.get("business_summary", "").splitlines():
        line = line.strip()
        if line.startswith("•") or line.startswith("-"):
            parts.append(f"<ul><li>{_md_bold(line.lstrip('•-').strip())}</li></ul>")
        elif line.startswith("**") and line.endswith("**"):
            parts.append(f"<h4>{line.strip('*')}</h4>")
        elif line:
            parts.append(f"<p>{_md_bold(line)}</p>")
    parts.append("<h3>Technical Summary</h3>")
    for line in exec_summary.get("technical_summary", "").splitlines():
        line = line.strip()
        if line.startswith("•") or line.startswith("-"):
            parts.append(f"<ul><li>{_md_bold(line.lstrip('•-').strip())}</li></ul>")
        elif line.startswith("**") and line.endswith("**"):
            parts.append(f"<h4>{line.strip('*')}</h4>")
        elif line:
            parts.append(f"<p>{_md_bold(line)}</p>")
    out.append(("## Executive Summary", "\n".join(parts)))

    # ── Job Architecture ─────────────────────────────────────────────────────
    comp_types = sorted({c.get("component_type","") for c in components})
    parts = [_kpi_grid(
        ("Total Components", len(components)),
        ("Unique Types", len(comp_types)),
    )]
    parts.append(_diagram_box("Pipeline Architecture — High Level", pipeline_svg))
    parts.append("<h3>Component Types</h3>")
    parts.append(_ul(comp_types))
    out.append(("## Job Architecture", "\n".join(parts)))

    # ── Source Architecture ──────────────────────────────────────────────────
    src_names = inv.get("source_names", [])
    src_schemas = inv.get("sources", [])
    parts = [f"<p><strong>Source Systems:</strong> {', '.join(src_names) or 'None detected'}</p>"]
    if src_schemas:
        rows = [[s.get("name",""), s.get("component_type",""), s.get("schema_name",""),
                 str(len(s.get("columns",[]))) + " cols"] for s in src_schemas]
        parts.append(_table_html(["Name", "Type", "Schema", "Columns"], rows))
    out.append(("## Source Architecture", "\n".join(parts)))

    # ── Target Architecture ──────────────────────────────────────────────────
    tgt_names = inv.get("target_names", [])
    tgt_schemas = inv.get("targets", [])
    parts = [f"<p><strong>Target Systems:</strong> {', '.join(tgt_names) or 'None detected'}</p>"]
    if tgt_schemas:
        rows = [[t.get("name",""), t.get("component_type",""), t.get("schema_name",""),
                 str(len(t.get("columns",[]))) + " cols"] for t in tgt_schemas]
        parts.append(_table_html(["Name", "Type", "Schema", "Columns"], rows))
    out.append(("## Target Architecture", "\n".join(parts)))

    # ── Source-To-Target Mapping ─────────────────────────────────────────────
    mappings = job_data.get("column_mappings", [])
    if mappings:
        rows = [[
            f'{m.get("Source Component","")}.{m.get("Source Column","")}',
            str(m.get("Expression",""))[:80],
            f'{m.get("Target Component","")}.{m.get("Target Column","")}',
            m.get("Data Type Conversion",""),
            m.get("Default Value",""),
        ] for m in mappings]
        section_html = _table_html(
            ["Source Column","Transformation","Target Column","Type Conversion","Default"], rows)
    else:
        section_html = "<p>No tMap column mappings found.</p>"
    out.append(("## Source-To-Target Mapping", section_html))

    # ── Transformation Architecture ──────────────────────────────────────────
    out.append(("## Transformation Architecture", _ul(sections["transformation"]["findings"])))

    # ── Job Flow Architecture ────────────────────────────────────────────────
    parts = [_ul(sections["job_flow"]["findings"])]
    parts.append(_diagram_box("Job Flow — Component & Subjob Diagram", flow_svg))
    out.append(("## Job Flow Architecture", "\n".join(parts)))

    # ── Column Lineage ───────────────────────────────────────────────────────
    lineage_rows = sections["column_lineage"].get("lineage_rows", [])
    if lineage_rows:
        rows = [[r.get("Source",""), str(r.get("Transformation",""))[:80], r.get("Target","")]
                for r in lineage_rows]
        section_html = _table_html(["Source","Transformation","Target"], rows)
    else:
        section_html = "<p>No column lineage data available.</p>"
    out.append(("## Column Lineage", section_html))

    # ── Validation / Error Handling / Audit / Performance / Security ────────
    out.append(("## Validation", _ul(sections["validation"]["findings"])))
    out.append(("## Error Handling", _ul(sections["error_handling"]["findings"])))
    out.append(("## Audit & Monitoring", _ul(sections["audit_monitoring"]["findings"])))
    out.append(("## Performance", _ul(sections["performance"]["findings"])))
    out.append(("## Security", _ul(sections["security"]["findings"])))

    # ── Dependency Architecture ──────────────────────────────────────────────
    parts = [_ul(sections["dependency"]["findings"])]
    parts.append(_diagram_box("Job Dependency Diagram", dep_svg))
    out.append(("## Dependency Architecture", "\n".join(parts)))

    # ── Testing ──────────────────────────────────────────────────────────────
    parts = [_kpi_grid(
        ("Unit Tests",            len(testing.get("unit_tests", []))),
        ("Validation SQL",        len(testing.get("validation_sql", []))),
        ("Reconciliation Rules",  len(testing.get("reconciliation_rules", []))),
        ("Src vs Tgt Checks",     len(testing.get("src_vs_tgt", []))),
    )]
    if testing.get("unit_tests"):
        parts.append("<h3>Unit Tests</h3>")
        rows = [[t.get("test_name",""), t.get("component",""), t.get("test_type",""), t.get("description","")]
                for t in testing["unit_tests"][:20]]
        parts.append(_table_html(["Test Name","Component","Type","Description"], rows))
    out.append(("## Testing", "\n".join(parts)))

    # ── Migration Assessment ──────────────────────────────────────────────────
    cr = assessment.get("cloud_readiness", {})
    rag = cr.get("rag","AMBER")
    parts = [_kpi_grid(
        ("Cloud Readiness",       cr.get("readiness","N/A")),
        ("Estimated Effort (hrs)",assessment.get("effort_estimation",{}).get("estimated_hours","—")),
        ("Unsupported Components",len(assessment.get("unsupported_components",[]))),
        ("Migration Risks",       len(assessment.get("migration_risks",[]))),
    )]
    parts.append(f"<p>RAG Status: {_rag_badge(rag)}</p>")
    if assessment.get("migration_risks"):
        rows = [[r.get("component",""), r.get("risk",""), r.get("impact",""), r.get("recommendation","")]
                for r in assessment["migration_risks"]]
        parts.append("<h3>Risk Register</h3>")
        parts.append(_table_html(["Component","Risk","Impact","Recommendation"], rows))
    parts.append("<h3>Recommendations</h3>")
    parts.append(_ul(assessment.get("recommendations",[])))
    out.append(("## Migration Assessment", "\n".join(parts)))

    # ── AI Executive Summary ──────────────────────────────────────────────────
    parts = ["<h3>Risks</h3>"]
    risks_text = exec_summary.get("risks","")
    if isinstance(risks_text, list):
        parts.append(_ul(risks_text))
    else:
        parts.append(_ul([l.lstrip("-•").strip() for l in risks_text.splitlines() if l.strip()]))
    parts.append("<h3>Opportunities</h3>")
    opp_text = exec_summary.get("opportunities","")
    for line in (opp_text if isinstance(opp_text, list) else opp_text.splitlines()):
        line = str(line).strip()
        if line:
            parts.append(f"<p>{_md_bold(line)}</p>")
    parts.append("<h3>Recommendations</h3>")
    recs_text = exec_summary.get("recommendations","")
    if isinstance(recs_text, list):
        parts.append(_ul(recs_text))
    else:
        parts.append(_ul([l.lstrip("-•").strip() for l in recs_text.splitlines() if l.strip()]))
    out.append(("## AI Executive Summary", "\n".join(parts)))

    return out


def build_tdd_html(job_data: dict, selected_headings: list[str] | None = None) -> str:
    """Build a rich HTML TDD with embedded SVG diagrams, KPI grids, and tables.

    `selected_headings`, if given, restricts the document to only the matching
    `## Heading` entries from build_tdd_sections_html (same heading strings used
    by build_tdd_markdown's section list) — letting the TDD section picker
    filter the export without falling back to the plain-text markdown path
    and losing styling/diagrams.
    """
    job_name = job_data.get("job_name", "Unknown")
    all_sections = build_tdd_sections_html(job_data)
    if selected_headings is not None:
        wanted = set(selected_headings)
        all_sections = [(h, c) for h, c in all_sections if h in wanted]

    body_parts = [f"<h1>Technical Design Document — {job_name}</h1>"]
    for heading, content_html in all_sections:
        title = heading.lstrip("#").strip()
        body_parts.append(f"<h2>{title}</h2>")
        body_parts.append(content_html)

    body_html = "\n".join(body_parts)
    return f"""<!doctype html>
<html>
<head>
  <meta charset='utf-8'/>
  <title>{job_name} — Technical Design Document</title>
  {_HTML_STYLE}
</head>
<body>
{body_html}
</body>
</html>"""




# ── Markdown (unchanged, text-only) ──────────────────────────────────────────

def build_tdd_markdown(job_data: dict) -> str:
    """Assemble the full TDD markdown for a single job."""
    job_name   = job_data.get("job_name", "Unknown")
    components = job_data.get("components", [])
    inv        = build_source_target_inventory(job_data)

    sections     = tdd_sections.generate_all_sections(job_data)
    testing      = build_testing_architecture(job_data)
    assessment   = build_migration_assessment(job_data)
    exec_summary = build_executive_summary(job_data)

    col_lineage = sections["column_lineage"]

    lines = [
        f"# Technical Design Document — {job_name}", "",
        "## Executive Summary",
        "### Business Summary", exec_summary["business_summary"], "",
        "### Technical Summary", exec_summary["technical_summary"], "",
        "## Job Architecture",
        f"- Component Count: {len(components)}",
        f"- Component Types: {', '.join(sorted({c.get('component_type','') for c in components})) or 'N/A'}", "",
        "## Source Architecture",
        f"- Source Systems: {', '.join(inv.get('source_names', [])) or 'None detected'}", "",
        "## Target Architecture",
        f"- Target Systems: {', '.join(inv.get('target_names', [])) or 'None detected'}", "",
        "## Source-To-Target Mapping",
    ] + ([
        "| Source Column | Transformation | Target Column | Data Type Conversion | Default Value |",
        "|---|---|---|---|---|",
    ] + [
        f'| {m.get("Source Component","")}.{m.get("Source Column","")} '
        f'| {str(m.get("Expression",""))[:80]} '
        f'| {m.get("Target Component","")}.{m.get("Target Column","")} '
        f'| {m.get("Data Type Conversion","")} '
        f'| {m.get("Default Value","")} |'
        for m in job_data.get("column_mappings", [])
    ] if job_data.get("column_mappings") else ["- No tMap column mappings found."]) + [
        "",
        "## Transformation Architecture", _bullets(sections["transformation"]["findings"]), "",
        "## Job Flow Architecture",        _bullets(sections["job_flow"]["findings"]), "",
        "## Column Lineage",
    ] + ([
        "| Source | Transformation | Target |", "|---|---|---|",
    ] + [
        f'| {r["Source"]} | {str(r["Transformation"])[:80]} | {r["Target"]} |'
        for r in col_lineage.get("lineage_rows", [])
    ] if col_lineage.get("lineage_rows") else ["- No column lineage data available."]) + [
        "",
        "## Validation",             _bullets(sections["validation"]["findings"]),        "",
        "## Error Handling",         _bullets(sections["error_handling"]["findings"]),    "",
        "## Audit & Monitoring",     _bullets(sections["audit_monitoring"]["findings"]),  "",
        "## Performance",            _bullets(sections["performance"]["findings"]),       "",
        "## Security",               _bullets(sections["security"]["findings"]),          "",
        "## Dependency Architecture",_bullets(sections["dependency"]["findings"]),        "",
        "## Testing",
        f"- Unit Tests: {len(testing['unit_tests'])}",
        f"- Validation SQL Checks: {len(testing['validation_sql'])}",
        f"- Reconciliation Rules: {len(testing['reconciliation_rules'])}",
        f"- Source vs Target Checks: {len(testing['src_vs_tgt'])}", "",
        "## Migration Assessment",
        f"- Cloud Readiness: {assessment['cloud_readiness']['readiness']} (RAG: {assessment['cloud_readiness']['rag']})",
        f"- Unsupported Components: {len(assessment['unsupported_components'])}",
        f"- Migration Risks: {len(assessment['migration_risks'])}",
        f"- Estimated Effort: {assessment['effort_estimation']['estimated_hours']} hours",
        "### Recommendations", _bullets(assessment["recommendations"]), "",
        "## AI Executive Summary",
        "### Risks",        exec_summary["risks"],        "",
        "### Opportunities",exec_summary["opportunities"],"",
        "### Recommendations",exec_summary["recommendations"], "",
    ]
    return "\n".join(lines)


# ── PDF rendering ─────────────────────────────────────────────────────────────

def _write_tdd_pdf(path: str, markdown: str, title: str,
                   diagram_pngs: dict[str, bytes] | None = None) -> str:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                     Table, TableStyle, ListFlowable, ListItem, Image)

    doc = SimpleDocTemplate(path, pagesize=A4,
                            rightMargin=42, leftMargin=42, topMargin=40, bottomMargin=40,
                            title=title)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("TmaH1",  parent=styles["Heading1"], fontSize=18,
                              textColor=colors.HexColor("#0F172A")))
    styles.add(ParagraphStyle("TmaH2",  parent=styles["Heading2"], fontSize=13,
                              textColor=colors.HexColor("#1E293B"), spaceBefore=12))
    styles.add(ParagraphStyle("TmaH3",  parent=styles["Heading3"], fontSize=11,
                              textColor=colors.HexColor("#334155"), spaceBefore=8))
    styles.add(ParagraphStyle("TmaBody",parent=styles["BodyText"], fontSize=9.5, leading=13))

    story = []
    diagram_pngs = diagram_pngs or {}
    table_buffer: list[list[str]] = []
    bullet_buffer: list[str] = []

    def flush_table():
        nonlocal table_buffer
        if not table_buffer: return
        ncols = max(len(r) for r in table_buffer)
        rows  = [r + [""] * (ncols - len(r)) for r in table_buffer]
        cw    = (A4[0] - 84) / ncols
        tbl   = Table(rows, colWidths=[cw]*ncols, repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#0F172A")),
            ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
            ("FONTSIZE",   (0,0), (-1,-1), 8),
            ("GRID",       (0,0), (-1,-1), 0.5, colors.HexColor("#CBD5E1")),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white, colors.HexColor("#F8FAFC")]),
            ("VALIGN",     (0,0), (-1,-1), "MIDDLE"),
        ]))
        story.append(tbl); story.append(Spacer(1,8))
        table_buffer = []

    def flush_bullets():
        nonlocal bullet_buffer
        if not bullet_buffer: return
        items = [ListItem(Paragraph(b, styles["TmaBody"])) for b in bullet_buffer]
        story.append(ListFlowable(items, bulletType="bullet", start="circle"))
        story.append(Spacer(1,6))
        bullet_buffer = []

    # Diagram insertion markers
    _DIAG_RE = re.compile(r"^\[DIAGRAM:(.+)\]$")

    for line in markdown.splitlines():
        stripped = line.strip()
        m = _DIAG_RE.match(stripped)
        if m:
            flush_table(); flush_bullets()
            key = m.group(1)
            if key in diagram_pngs:
                img_io = io.BytesIO(diagram_pngs[key])
                img = Image(img_io, width=A4[0]-84, height=200, kind="proportional")
                story.append(img)
                story.append(Spacer(1,8))
            continue

        if stripped.startswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if set("".join(cells).replace("-","")) != set():
                table_buffer.append(cells)
            continue
        else:
            flush_table()

        if stripped.startswith("- "):
            bullet_buffer.append(stripped[2:]); continue
        else:
            flush_bullets()

        if   stripped.startswith("# "):  story.append(Paragraph(stripped[2:], styles["TmaH1"]))
        elif stripped.startswith("## "): story.append(Paragraph(stripped[3:], styles["TmaH2"]))
        elif stripped.startswith("### "):story.append(Paragraph(stripped[4:], styles["TmaH3"]))
        elif stripped.startswith("```"): continue
        elif stripped: story.append(Paragraph(stripped, styles["TmaBody"]))
        else: story.append(Spacer(1,4))

    flush_table(); flush_bullets()
    doc.build(story)
    return path


# ── Export Entry Points ────────────────────────────────────────────────────────

def export_tdd(job_data: dict, output_dir: str) -> dict[str, str]:
    """Export the full TDD for one job to all 4 formats. Returns a dict of
    format -> file path. Filenames follow <Job_Name>_TDD.<ext>."""
    os.makedirs(output_dir, exist_ok=True)
    job_name = _safe_job_name(job_data.get("job_name", "Unknown"))
    title    = f"{job_data.get('job_name', 'Unknown')} — Technical Design Document"

    # ── 1. HTML — rich with embedded SVG diagrams ─────────────────────────────
    html_path = os.path.join(output_dir, f"{job_name}_TDD.html")
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(build_tdd_html(job_data))

    # ── 2. Markdown — plain text ──────────────────────────────────────────────
    markdown  = build_tdd_markdown(job_data)
    md_path   = os.path.join(output_dir, f"{job_name}_TDD.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(markdown)

    # ── 3. DOCX — markdown with PNG diagrams inserted ────────────────────────
    docx_path = os.path.join(output_dir, f"{job_name}_TDD.docx")
    write_docx(docx_path, title, markdown)
    # Attempt to inject PNG diagrams into DOCX after the fact
    try:
        _inject_diagrams_docx(docx_path, job_data)
    except Exception:
        pass  # fallback to text-only DOCX

    # ── 4. PDF — markdown with PNG diagrams ──────────────────────────────────
    pdf_path = os.path.join(output_dir, f"{job_name}_TDD.pdf")
    # Build PNG diagram bytes — pure-Python, no Graphviz needed
    from app.tiap.documentation import diagram_renderer as _dr_pdf
    diagram_pngs: dict[str, bytes] = {}
    _sc_pdf = {s: 0 for s in _PIPELINE_STAGES}
    for _c_pdf in job_data.get("components", []):
        _sc_pdf[_classify_stage(_c_pdf.get("component_type", ""))] += 1
    _p = _dr_pdf.build_pipeline_png(job_data, _PIPELINE_STAGES, _sc_pdf)
    if _p: diagram_pngs["pipeline"] = _p
    _f = _dr_pdf.build_flow_png(job_data, _subjob_list2, _type_lookup2, _data_edges2, _trigger_edges2)
    if _f: diagram_pngs["flow"] = _f
    _d = _dr_pdf.build_dependency_png(_job_name2, _child_jobs2)
    if _d: diagram_pngs["dep"] = _d

    # Build markdown with diagram markers for PDF
    md_with_markers = markdown
    md_with_markers = md_with_markers.replace(
        "## Job Architecture\n",
        "## Job Architecture\n[DIAGRAM:pipeline]\n")
    md_with_markers = md_with_markers.replace(
        "## Job Flow Architecture\n",
        "## Job Flow Architecture\n[DIAGRAM:flow]\n")
    md_with_markers = md_with_markers.replace(
        "## Dependency Architecture\n",
        "## Dependency Architecture\n[DIAGRAM:dep]\n")
    _write_tdd_pdf(pdf_path, md_with_markers, title, diagram_pngs)

    return {
        "html":     html_path,
        "markdown": md_path,
        "docx":     docx_path,
        "pdf":      pdf_path,
    }


def _inject_diagrams_docx(docx_path: str, job_data: dict) -> None:
    """Insert PNG diagrams into the DOCX after relevant section headings."""
    from docx import Document
    from docx.shared import Inches
    doc = Document(docx_path)
    new_paragraphs = []
    for para in doc.paragraphs:
        new_paragraphs.append(para)
        heading_text = para.text.strip()
        dot_fn = None
        if heading_text == "Job Architecture":
            dot_fn = _build_pipeline_dot
        elif heading_text == "Job Flow Architecture":
            dot_fn = _build_flow_dot
        elif heading_text == "Dependency Architecture":
            dot_fn = _build_dependency_dot
        if dot_fn:
            try:
                from app.tiap.documentation import diagram_renderer as _dr_docx
                import re as _re_docx
                _sc_d = {s: 0 for s in _PIPELINE_STAGES}
                for _cc in job_data.get("components", []):
                    _sc_d[_classify_stage(_cc.get("component_type", ""))] += 1
                if dot_fn is _build_pipeline_dot:
                    png_bytes = _dr_docx.build_pipeline_png(job_data, _PIPELINE_STAGES, _sc_d)
                elif dot_fn is _build_flow_dot:
                    _TREX = _re_docx.compile(r"SUBJOB|RUN_IF|COMPONENT_OK|COMPONENT_ERROR|ITERATE|^OK$|^ERROR$", _re_docx.I)
                    _cps = job_data.get("components", [])
                    _cns = job_data.get("connections", [])
                    _par = {(_c.get("unique_name") or _c.get("component_type", f"c{_i}")): None for _i, _c in enumerate(_cps)}
                    def _fn(x):
                        while _par[x]: x = _par[x]
                        return x
                    def _un(a, b):
                        ra, rb = _fn(a), _fn(b)
                        if ra != rb: _par[ra] = rb
                    _de, _te = [], []
                    for _cn in _cns:
                        _ss = _cn.get("source",""); _tt = _cn.get("target",""); _ctt = _cn.get("connector_type","")
                        if _ss in _par and _tt in _par:
                            (_te if _TREX.search(_ctt or "") else _de).append((_ss,_tt,_ctt))
                            if not _TREX.search(_ctt or ""): _un(_ss,_tt)
                    _gr = {}
                    for _nm in _par:
                        _gr.setdefault(_fn(_nm), []).append(_nm)
                    _tl = {(_c.get("unique_name") or _c.get("component_type")): _c.get("component_type","") for _c in _cps}
                    png_bytes = _dr_docx.build_flow_png(job_data, list(_gr.values()), _tl, _de, _te)
                else:
                    _jn = job_data.get("job_name","Unknown")
                    _cj = [_c.get("unique_name","") for _c in job_data.get("components",[]) if _c.get("component_type")=="tRunJob"]
                    png_bytes = _dr_docx.build_dependency_png(_jn, _cj)
                if png_bytes:
                    img_io = io.BytesIO(png_bytes)
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(img_io, width=Inches(6))
                    # Move the new paragraph after current
                    para._element.addnext(p._element)
            except Exception:
                pass
    doc.save(docx_path)
