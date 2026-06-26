import os
import shutil
from typing import Dict, Mapping

from docx import Document


TEMPLATE_DIR = "templates"
DEFAULT_TEMPLATE_PATH = os.path.join(TEMPLATE_DIR, "default_template.docx")
CUSTOM_TEMPLATE_PATH = os.path.join(TEMPLATE_DIR, "custom_template.docx")
TEMPLATE_SESSION_KEY = "report_pack_template_path"

PLACEHOLDERS = {
    "EXECUTIVE_SUMMARY": "{{EXECUTIVE_SUMMARY}}",
    "REPOSITORY_OVERVIEW": "{{REPOSITORY_OVERVIEW}}",
    "READINESS_SCORE": "{{READINESS_SCORE}}",
    "FLOWCHARTS": "{{FLOWCHARTS}}",
    "TECH_DOC": "{{TECH_DOC}}",
    "FUNC_DOC": "{{FUNC_DOC}}",
    "KT_DOC": "{{KT_DOC}}",
    "MIGRATION": "{{MIGRATION}}",
    "TEST_CASES": "{{TEST_CASES}}",
    "RECOMMENDATIONS": "{{RECOMMENDATIONS}}",
}


def ensure_template_dir() -> str:
    os.makedirs(TEMPLATE_DIR, exist_ok=True)
    return TEMPLATE_DIR


def default_template_exists() -> bool:
    return os.path.exists(DEFAULT_TEMPLATE_PATH)


def save_custom_template(uploaded_file) -> str:
    ensure_template_dir()
    with open(CUSTOM_TEMPLATE_PATH, "wb") as handle:
        handle.write(uploaded_file.getbuffer())
    return CUSTOM_TEMPLATE_PATH


def restore_default_template() -> str:
    ensure_template_dir()
    return DEFAULT_TEMPLATE_PATH


def active_template_label(template_path: str) -> str:
    if not template_path:
        return "No template selected"
    if os.path.abspath(template_path) == os.path.abspath(DEFAULT_TEMPLATE_PATH):
        return "Default Template"
    if os.path.abspath(template_path) == os.path.abspath(CUSTOM_TEMPLATE_PATH):
        return "Custom Template"
    return os.path.basename(template_path)


def copy_default_template_from(source_path: str) -> str:
    ensure_template_dir()
    shutil.copyfile(source_path, DEFAULT_TEMPLATE_PATH)
    return DEFAULT_TEMPLATE_PATH


def build_placeholder_values(sections: Mapping[str, str], test_cases: str = "") -> Dict[str, str]:
    flowcharts = "\n\n".join([
        "# Technical Flowchart",
        sections.get("Technical Flowchart", ""),
        "# Business Flowchart",
        sections.get("Business Flowchart", ""),
        "# Repository Flowchart",
        sections.get("Repository Flowchart", ""),
    ])
    return {
        PLACEHOLDERS["EXECUTIVE_SUMMARY"]: sections.get("Executive Summary", ""),
        PLACEHOLDERS["REPOSITORY_OVERVIEW"]: sections.get("Repository Overview", ""),
        PLACEHOLDERS["READINESS_SCORE"]: sections.get("Readiness Scores", ""),
        PLACEHOLDERS["FLOWCHARTS"]: flowcharts,
        PLACEHOLDERS["TECH_DOC"]: sections.get("Technical Documentation", ""),
        PLACEHOLDERS["FUNC_DOC"]: sections.get("Functional Documentation", ""),
        PLACEHOLDERS["KT_DOC"]: sections.get("KT Documentation", ""),
        PLACEHOLDERS["MIGRATION"]: sections.get("Migration Assessment", ""),
        PLACEHOLDERS["TEST_CASES"]: test_cases or "Test cases are not generated in this phase.",
        PLACEHOLDERS["RECOMMENDATIONS"]: sections.get("Recommendations", ""),
    }


def render_template_docx(
    template_path: str,
    output_path: str,
    placeholder_values: Mapping[str, str],
) -> Dict[str, int]:
    document = Document(template_path)
    replacement_count = 0

    for paragraph in document.paragraphs:
        replacement_count += _replace_in_paragraph(paragraph, placeholder_values)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replacement_count += _replace_in_paragraph(paragraph, placeholder_values)

    if replacement_count == 0:
        document.add_page_break()
        _add_heading(document, "Complete Assessment")
        for placeholder, content in placeholder_values.items():
            title = placeholder.strip("{}").replace("_", " ").title()
            _add_heading(document, title)
            document.add_paragraph(str(content or "No content generated."))

    document.save(output_path)
    return {"replacement_count": replacement_count}


def _replace_in_paragraph(paragraph, replacements: Mapping[str, str]) -> int:
    original = paragraph.text
    if not original:
        return 0

    updated = original
    count = 0
    for placeholder, value in replacements.items():
        if placeholder in updated:
            count += updated.count(placeholder)
            updated = updated.replace(placeholder, str(value or ""))

    if updated == original:
        return 0

    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = updated
    else:
        paragraph.add_run(updated)
    return count


def _add_heading(document: Document, text: str) -> None:
    try:
        document.add_heading(text, level=1)
    except KeyError:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
