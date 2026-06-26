import os
from datetime import datetime
from typing import Any, Dict, Iterable, List, Sequence

from docx import Document
from docx.shared import Pt

from app.tiap.documentation.executive_summary_generator import ExecutiveSummaryGenerator
from app.tiap.documentation.functional_doc_generator import FunctionalDocGenerator
from app.tiap.documentation.kt_doc_generator import KTDocGenerator
from app.tiap.documentation.migration_doc_generator import MigrationDocGenerator
from app.tiap.documentation.template_manager import (
    build_placeholder_values,
    render_template_docx,
)
from app.tiap.documentation.export_utils import markdown_fragment_to_html
from app.tiap.documentation.technical_doc_generator import TechnicalDocGenerator
from app.tiap.graph.flowchart_generator import FlowchartGenerator
from app.reports.excel_export import write_complete_assessment_excel
from app.reports.json_export import write_complete_assessment_json
from app.analyzers.routine_analyzer import analyze_routines
from app.analyzers.joblet_analyzer import analyze_joblets
from app.analyzers.java_risk_analyzer import analyze_java_risks
from app.tiap.inventory.inventory_parser import InventoryParser
from app.analyzers.readiness_scorer import RepositoryScoring
from app.analyzers.models import RepositoryOverview
from app.analyzers.version_upgrade_analyzer import UpgradePathAnalyzer
from app.analyzers.version_compatibility_engine import VersionCompatibilityEngine
from app.repository.repository_type_detector import RepositoryTypeDetector
from app.repository.enterprise_feature_detector import EnterpriseFeatureDetector
from app.risk_engine.risk_analyzer import RiskAnalyzer


REPORT_PACK_SESSION_KEY = "ai_report_pack"
REPORT_PACK_FILENAME = "Complete_Assessment.docx"


def build_report_pack(
    all_jobs: Sequence[Dict[str, Any]],
    repository_path: str = None,
    output_dir: str = "output",
    effort: Dict[str, Any] = None,
    auto_fix_recs: Sequence[Dict[str, Any]] = None,
    technical_template: str = None,
    report_template_path: str = None,
    test_cases: str = "",
    repository_ai_context: Dict[str, Any] = None,
) -> Dict[str, Any]:
    """Build all report sections and export the complete DOCX pack."""
    os.makedirs(output_dir, exist_ok=True)

    sections = build_report_pack_sections(
        all_jobs=all_jobs,
        repository_path=repository_path,
        effort=effort,
        auto_fix_recs=auto_fix_recs,
        technical_template=technical_template,
        repository_ai_context=repository_ai_context,
    )
    docx_path = os.path.join(output_dir, REPORT_PACK_FILENAME)
    html_path = os.path.join(output_dir, "Complete_Assessment.html")
    pdf_path  = os.path.join(output_dir, "Complete_Assessment.pdf")
    excel_path = os.path.join(output_dir, "Complete_Assessment.xlsx")
    json_path  = os.path.join(output_dir, "Complete_Assessment.json")

    template_result = {}
    if report_template_path and os.path.exists(report_template_path):
        placeholder_values = build_placeholder_values(sections, test_cases=test_cases)
        template_result = render_template_docx(report_template_path, docx_path, placeholder_values)
    else:
        write_complete_assessment_docx(docx_path, sections)

    # Always generate HTML, PDF, Excel, and JSON from sections (independent of template)
    write_complete_assessment_html(html_path, sections)
    write_complete_assessment_pdf(pdf_path, sections)
    write_complete_assessment_excel(excel_path, sections, all_jobs, repository_path)

    overview_model = _build_repository_overview_model(all_jobs, repository_path)
    write_complete_assessment_json(
        json_path, sections, all_jobs, repository_path,
        repository_overview=overview_model.to_dict(),
    )

    return {
        "sections": sections,
        "docx_path": docx_path,
        "html_path": html_path,
        "pdf_path": pdf_path,
        "excel_path": excel_path,
        "json_path": json_path,
        "template_path": report_template_path,
        "template_result": template_result,
        "generated_at": datetime.now().isoformat(timespec="seconds"),
    }



def _routine_assessment_md(all_jobs: Sequence[Dict[str, Any]]) -> str:
    try:
        data = analyze_routines(all_jobs)
        routines = data.get("routines", [])
        if not routines:
            return "No custom routines detected."
        lines = ["| Routine Name | Usage Count | Complexity | Migration Risk |",
                 "| --- | --- | --- | --- |"]
        for r in routines:
            lines.append(f"| {r.get('name','?')} | {r.get('usage_count', r.get('job_count','?'))} | {r.get('complexity','?')} | {r.get('risk_level', r.get('migration_risk','?'))} |")
        return "\n".join(lines)
    except Exception as e:
        return f"Routine assessment unavailable: {e}"


def _joblet_assessment_md(all_jobs: Sequence[Dict[str, Any]]) -> str:
    try:
        data = analyze_joblets(all_jobs)
        joblets = data.get("joblets", [])
        if not joblets:
            return "No joblets detected."
        lines = ["| Joblet Name | Used In Jobs | Complexity | Action Required |",
                 "| --- | --- | --- | --- |"]
        for j in joblets:
            jobs_list = ", ".join(j.get("jobs_using", []))
            lines.append(f"| {j.get('name','?')} | {jobs_list or j.get('job_count','?')} | {j.get('complexity','?')} | {j.get('risk_level','Review')} |")
        return "\n".join(lines)
    except Exception as e:
        return f"Joblet assessment unavailable: {e}"


def _java_risk_md(all_jobs: Sequence[Dict[str, Any]]) -> str:
    try:
        data = analyze_java_risks(all_jobs)
        score = data.get("java_risk_score", data.get("overall_risk", "N/A"))
        jobs = data.get("job_risks", data.get("jobs", []))
        lines = [f"**Java Risk Score:** {score}\n",
                 "| Job | Risk Level | Issues Detected |",
                 "| --- | --- | --- |"]
        for j in jobs:
            issues = ", ".join(j.get("risks", j.get("issues", [])))
            lines.append(f"| {j.get('job_name', j.get('name','?'))} | {j.get('risk_level','?')} | {issues or 'None'} |")
        return "\n".join(lines)
    except Exception as e:
        return f"Java risk assessment unavailable: {e}"


def _doc_readiness_md(all_jobs: Sequence[Dict[str, Any]], repository_path: str = None) -> str:
    try:
        scores = RepositoryScoring().score(all_jobs)
        overall = scores.get("overall_score", scores.get("readiness_score", "N/A"))
        lines = [f"**Overall Documentation Readiness Score:** {overall}\n",
                 "| Category | Score |", "| --- | --- |"]
        for k, v in scores.items():
            if k != "overall_score":
                lines.append(f"| {k.replace('_', ' ').title()} | {v} |")
        return "\n".join(lines)
    except Exception as e:
        return f"Documentation readiness unavailable: {e}"


def _appendix_md(all_jobs: Sequence[Dict[str, Any]]) -> str:
    import json, os
    sections = []

    def _json_table(path, key=None):
        if not os.path.exists(path):
            return ""
        try:
            data = json.loads(open(path).read())
            if key:
                data = data.get(key, data)
            if isinstance(data, list) and data and isinstance(data[0], dict):
                headers = list(data[0].keys())
                rows = [["| " + " | ".join(headers) + " |"],
                        ["| " + " | ".join(["---"] * len(headers)) + " |"]]
                for row in data:
                    rows.append(["| " + " | ".join(str(row.get(h, "")) for h in headers) + " |"])
                return "\n".join(r[0] for r in rows)
        except Exception:
            pass
        return ""

    sections.append("## Component Inventory\n" + (_json_table("output/repository_inventory.json") or "Not available"))
    ctx_lines = []
    for job in all_jobs[:20]:
        jd = job.get("job_data", job)
        for ctx in jd.get("context_variables", []):
            ctx_lines.append(f"| {jd.get('job_name','?')} | {ctx.get('name','?')} | {ctx.get('value','?')} |")
    if ctx_lines:
        sections.append("## Context Variables\n| Job | Variable | Value |\n| --- | --- | --- |\n" + "\n".join(ctx_lines))
    else:
        sections.append("## Context Variables\nNone detected.")
    sections.append("## Job Dependencies\n" + (_json_table("output/dependency_summary.json") or "Not available"))
    sections.append("## Joblets\n" + _joblet_assessment_md(all_jobs))
    sections.append("## Routines\n" + _routine_assessment_md(all_jobs))
    dep_data = json.loads(open("output/repository_inventory.json").read()) if os.path.exists("output/repository_inventory.json") else {}
    deprecated = dep_data.get("deprecated_components", [])
    if deprecated:
        dep_lines = ["| Component | Job | Replacement |", "| --- | --- | --- |"]
        for d in deprecated:
            dep_lines.append(f"| {d.get('component',d.get('name','?'))} | {d.get('job','?')} | {d.get('replacement','N/A')} |")
        sections.append("## Deprecated Components\n" + "\n".join(dep_lines))
    else:
        sections.append("## Deprecated Components\nNone detected.")
    sections.append("## Repository Statistics\n" + (_json_table("output/repository_summary.json") or "Not available"))
    return "\n\n".join(sections)


def _build_repository_overview_model(
    all_jobs: Sequence[Dict[str, Any]],
    repository_path: str = None,
) -> RepositoryOverview:
    """Build a RepositoryOverview model from jobs and repository metadata."""
    inventory = InventoryParser().build_inventory(all_jobs, repository_path)
    scoring = RepositoryScoring().score(all_jobs)

    repo_type_info: Dict[str, Any] = {"type": "Unknown", "source_version": "UNKNOWN"}
    if repository_path:
        try:
            detector = RepositoryTypeDetector()
            repo_type_info = {
                "type": detector.detect_from_path(repository_path).get("type", "Unknown"),
                "source_version": detector.extract_source_version_from_path(repository_path),
            }
        except Exception:
            pass

    enterprise_features_info: Dict[str, Any] = {}
    try:
        enterprise_features_info = EnterpriseFeatureDetector().detect_from_jobs(all_jobs)
    except Exception:
        pass

    source_version = repo_type_info.get("source_version", "UNKNOWN")
    target_version = "Talend 8"

    target_versions: List[str] = []
    try:
        if source_version not in (None, "UNKNOWN"):
            target_versions = [
                t["version"]
                for t in VersionCompatibilityEngine().get_supported_targets(source_version)
            ]
    except Exception:
        pass

    _risk_rank = {"LOW": 0, "MEDIUM": 1, "HIGH": 2, "CRITICAL": 3}
    repo_risk = "LOW"
    total_findings = 0
    total_blockers = 0
    for job_entry in all_jobs:
        job = job_entry.get("job_data", job_entry) if isinstance(job_entry, dict) else job_entry
        try:
            for r in RiskAnalyzer().analyze(job):
                if _risk_rank.get(r.get("risk", "LOW"), 0) > _risk_rank.get(repo_risk, 0):
                    repo_risk = r.get("risk", "LOW")
        except Exception:
            pass
        if source_version not in (None, "UNKNOWN"):
            try:
                path_result = UpgradePathAnalyzer().analyze_job(job, source_version, target_version)
                total_findings += len(path_result.get("componentFindings", []))
                total_blockers += len(path_result.get("blockers", []))
            except Exception:
                pass

    if source_version in (None, "UNKNOWN"):
        upgrade_path_summary = "Source version unknown — upgrade path could not be determined."
    elif total_blockers:
        upgrade_path_summary = f"{total_blockers} job(s) blocked from {source_version} to {target_version}."
    elif total_findings:
        upgrade_path_summary = f"{total_findings} component change(s) required to upgrade from {source_version} to {target_version}."
    else:
        upgrade_path_summary = f"Clean upgrade path available from {source_version} to {target_version}."

    upgrade_path_info = {
        "targetVersions": target_versions,
        "migrationRisk": repo_risk,
        "summary": upgrade_path_summary,
    }

    return RepositoryOverview.from_inventory_and_scoring(
        inventory, scoring, repo_type_info, enterprise_features_info, upgrade_path_info
    )


def _repository_overview_section(overview: RepositoryOverview) -> str:
    """Render RepositoryOverview as a markdown section for export."""
    d = overview.to_dict()
    lines = [
        "# Repository Overview",
        "",
        f"- Repository Type: {d.get('repositoryType', 'Unknown')}",
        f"- Source Version: {d.get('sourceVersion', 'UNKNOWN')}",
        f"- Total Jobs: {d.get('totalJobs', 0)}",
        f"- Total Components: {d.get('totalComponents', 0)}",
        f"- Total Routines: {d.get('totalRoutines', 0)}",
        f"- Total Joblets: {d.get('totalJoblets', 0)}",
        f"- Complexity Score: {d.get('complexityScore', 0)}%",
        f"- Migration Readiness Score: {d.get('migrationReadinessScore', 0)}%",
        f"- Cloud Readiness Score: {d.get('cloudReadinessScore', 0)}%",
        f"- Testing Readiness Score: {d.get('testingReadinessScore', 0)}%",
        f"- Migration Risk: {d.get('migrationRisk', 'LOW')}",
        f"- Upgrade Path Summary: {d.get('upgradePathSummary', '')}",
    ]
    target_versions = d.get("targetVersions", [])
    if target_versions:
        lines.append(f"- Supported Target Versions: {', '.join(target_versions)}")
    enterprise_features = d.get("enterpriseFeatures", [])
    if enterprise_features:
        lines.append("")
        lines.append("## Enterprise Features Detected")
        for feat in enterprise_features:
            lines.append(f"- {feat}")
    return "\n".join(lines)


def _upgrade_path_section(
    all_jobs: Sequence[Dict[str, Any]],
    repository_path: str = None,
) -> str:
    """Build the Upgrade Path section with per-job component findings."""
    source_version: str = "UNKNOWN"

    # 1. Try to detect from repository path
    if repository_path:
        try:
            source_version = RepositoryTypeDetector().extract_source_version_from_path(repository_path)
        except Exception:
            pass

    # 2. Fall back to source_version on job data
    if source_version in (None, "UNKNOWN"):
        for job_entry in all_jobs:
            job = job_entry.get("job_data", job_entry) if isinstance(job_entry, dict) else job_entry
            sv = job.get("source_version")
            if sv and sv not in (None, "UNKNOWN"):
                source_version = sv
                break

    target_version = "Talend 8"

    if source_version in (None, "UNKNOWN"):
        return (
            "# Upgrade Path\n\n"
            "Source version could not be detected — upgrade path analysis unavailable."
        )

    lines = [
        "# Upgrade Path",
        "",
        f"- Source Version: {source_version}",
        f"- Target Version: {target_version}",
        "",
    ]

    try:
        path_report = UpgradePathAnalyzer().analyze_path(source_version, target_version)
        hops = path_report.get("hops", [])
        if hops:
            lines.append(f"- Upgrade Hops: {' → '.join(hops)}")
        renamed = path_report.get("renamedComponents", {})
        removed = path_report.get("removedComponents", [])
        param_changes = path_report.get("parameterChanges", {})
        if renamed:
            lines.extend(["", "## Renamed Components", "| Old Name | New Name |", "| --- | --- |"])
            for old, new in renamed.items():
                lines.append(f"| {old} | {new} |")
        if removed:
            lines.extend(["", "## Removed Components"])
            for comp in removed:
                lines.append(f"- {comp}")
        if param_changes:
            lines.extend(["", "## Parameter Changes"])
            for comp, changes in param_changes.items():
                lines.append(f"- {comp}: {changes}")
    except Exception as exc:
        lines.append(f"Upgrade path analysis error: {exc}")
        return "\n".join(lines)

    total_findings = 0
    total_blockers = 0
    job_rows: List[str] = []
    for job_entry in all_jobs:
        job = job_entry.get("job_data", job_entry) if isinstance(job_entry, dict) else job_entry
        job_name = job.get("job_name", "Unknown")
        try:
            result = UpgradePathAnalyzer().analyze_job(job, source_version, target_version)
            findings = result.get("componentFindings", [])
            blockers = result.get("blockers", [])
            total_findings += len(findings)
            total_blockers += len(blockers)
            status = "Blocked" if blockers else ("Changes Required" if findings else "Clean")
            job_rows.append(f"| {job_name} | {status} | {len(findings)} | {len(blockers)} |")
        except Exception:
            job_rows.append(f"| {job_name} | Error | — | — |")

    lines.extend([
        "",
        "## Per-Job Upgrade Summary",
        f"",
        f"- Total Component Findings: {total_findings}",
        f"- Total Blockers: {total_blockers}",
        "",
        "| Job Name | Status | Findings | Blockers |",
        "| --- | --- | --- | --- |",
    ])
    lines.extend(job_rows)
    return "\n".join(lines)


def _impact_lineage_section(all_jobs):
    """Documentation adapter for the canonical impact intelligence service."""
    try:
        from app.ui.impact_intelligence_dashboard import impact_report_markdown
        from app.impact_analysis.engine import ImpactLineageIntelligenceEngine
        return impact_report_markdown(ImpactLineageIntelligenceEngine().analyze(all_jobs))
    except Exception as exc:
        return f"Impact and lineage intelligence unavailable: {exc}"

def build_report_pack_sections(
    all_jobs: Sequence[Dict[str, Any]],
    repository_path: str = None,
    effort: Dict[str, Any] = None,
    auto_fix_recs: Sequence[Dict[str, Any]] = None,
    technical_template: str = None,
    repository_ai_context: Dict[str, Any] = None,
) -> Dict[str, str]:
    flows = FlowchartGenerator().generate(all_jobs)
    overview_model = _build_repository_overview_model(all_jobs, repository_path)
    sections = {
        "Executive Summary": ExecutiveSummaryGenerator().generate(all_jobs, repository_path, effort),
        "Repository Overview": _repository_overview_section(overview_model),
        "Upgrade Path": _upgrade_path_section(all_jobs, repository_path),
        "Readiness Scores": _readiness_scores(all_jobs, repository_path),
        "Impact & Lineage Intelligence": _impact_lineage_section(all_jobs),
        "Technical Flowchart": flows.get("technical_flow", "No technical flow detected"),
        "Business Flowchart": flows.get("business_flow", "No business flow detected"),
        "Repository Flowchart": flows.get("repository_flow", "Empty repository"),
        "Technical Documentation": TechnicalDocGenerator().generate(all_jobs, repository_path, technical_template),
        "Functional Documentation": FunctionalDocGenerator().generate(all_jobs),
        "KT Documentation": KTDocGenerator().generate(all_jobs),
        "Migration Assessment": MigrationDocGenerator().generate(all_jobs),
        "Recommendations": _recommendations(all_jobs, auto_fix_recs),
        "Routine Assessment": _routine_assessment_md(all_jobs),
        "Joblet Assessment": _joblet_assessment_md(all_jobs),
        "Java Risk Assessment": _java_risk_md(all_jobs),
        "Documentation Readiness": _doc_readiness_md(all_jobs, repository_path),
        "Test Cases": "",
        "Appendix": _appendix_md(all_jobs),
    }
    return _merge_repository_ai_context(sections, repository_ai_context or {})


def write_complete_assessment_docx(path: str, sections: Dict[str, str]) -> str:
    document = Document()
    document.core_properties.title = "Complete Assessment"
    document.core_properties.subject = "Talend Migration Accelerator AI Report Pack"

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    # --- Cover page ---
    document.add_heading("Complete Migration Assessment", level=0)
    document.add_paragraph("Generated by Artha Talend Migration Accelerator")
    document.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # --- Table of Contents ---
    document.add_page_break()
    document.add_heading("Table of Contents", level=1)
    for idx, title in enumerate(sections.keys(), start=1):
        document.add_paragraph(f"{idx}. {title}", style="List Number")

    # --- Sections ---
    for title, content in sections.items():
        document.add_page_break()
        document.add_heading(title, level=1)
        _add_markdown_content(document, content)

    document.save(path)
    return path


def write_complete_assessment_html(path: str, sections: Dict[str, str]) -> str:
    """Export all sections as a single styled HTML document."""
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    toc_items = "".join(
        f'<li><a href="#section-{idx}">{title}</a></li>'
        for idx, title in enumerate(sections.keys(), start=1)
    )
    body_sections = ""
    for idx, (title, content) in enumerate(sections.items(), start=1):
        html_content = markdown_fragment_to_html(str(content or ""))
        body_sections += f"""
        <div class="section" id="section-{idx}">
            <h1>{title}</h1>
            {html_content}
        </div>
        <div class="page-break"></div>
        """
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"/>
<title>Complete Migration Assessment</title>
<style>
  body {{ font-family: Arial, sans-serif; font-size: 11pt; color: #1a1a1a; margin: 40px; }}
  h1 {{ color: #1a3c6e; border-bottom: 2px solid #1a3c6e; padding-bottom: 6px; margin-top: 32px; }}
  h2 {{ color: #2563eb; margin-top: 24px; }}
  h3 {{ color: #374151; margin-top: 16px; }}
  ul {{ margin: 8px 0 8px 24px; }}
  li {{ margin-bottom: 4px; }}
  pre {{ background: #f3f4f6; padding: 12px; border-radius: 4px; font-size: 9pt; overflow-x: auto; }}
  .cover {{ text-align: center; padding: 80px 0 60px; }}
  .cover h1 {{ font-size: 28pt; border: none; }}
  .cover p {{ color: #6b7280; font-size: 11pt; }}
  .toc {{ background: #f9fafb; border: 1px solid #e5e7eb; padding: 24px 32px; border-radius: 6px; margin: 32px 0; }}
  .toc h2 {{ margin-top: 0; }}
  .toc ol {{ line-height: 2; }}
  .toc a {{ color: #2563eb; text-decoration: none; }}
  .section {{ margin-top: 48px; }}
  .page-break {{ page-break-after: always; }}
</style>
</head>
<body>
<div class="cover">
  <h1>Complete Migration Assessment</h1>
  <p>Generated by Artha Talend Migration Accelerator</p>
  <p>{generated_at}</p>
</div>
<div class="page-break"></div>
<div class="toc">
  <h2>Table of Contents</h2>
  <ol>{toc_items}</ol>
</div>
{body_sections}
</body>
</html>"""
    with open(path, "w", encoding="utf-8") as f:
        f.write(html)
    return path


def write_complete_assessment_pdf(pdf_path: str, sections: Dict[str, str]) -> str:
    """Export all sections as a single styled PDF using reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, PageBreak,
            HRFlowable,
        )

        doc = SimpleDocTemplate(
            pdf_path,
            pagesize=A4,
            leftMargin=20 * mm,
            rightMargin=20 * mm,
            topMargin=20 * mm,
            bottomMargin=20 * mm,
        )

        styles = getSampleStyleSheet()
        style_title  = ParagraphStyle("RPTitle",  parent=styles["Title"],   fontSize=24, textColor=colors.HexColor("#1a3c6e"), spaceAfter=12)
        style_sub    = ParagraphStyle("RPSub",    parent=styles["Normal"],  fontSize=11, textColor=colors.HexColor("#6b7280"), spaceAfter=6)
        style_h1     = ParagraphStyle("RPH1",     parent=styles["Heading1"], fontSize=16, textColor=colors.HexColor("#1a3c6e"), spaceBefore=18, spaceAfter=8)
        style_h2     = ParagraphStyle("RPH2",     parent=styles["Heading2"], fontSize=13, textColor=colors.HexColor("#2563eb"), spaceBefore=12, spaceAfter=6)
        style_h3     = ParagraphStyle("RPH3",     parent=styles["Heading3"], fontSize=11, textColor=colors.HexColor("#374151"), spaceBefore=8,  spaceAfter=4)
        style_body   = ParagraphStyle("RPBody",   parent=styles["Normal"],  fontSize=10, leading=14, spaceAfter=4)
        style_bullet = ParagraphStyle("RPBullet", parent=styles["Normal"],  fontSize=10, leading=14, leftIndent=16, spaceAfter=2)
        style_code   = ParagraphStyle("RPCode",   parent=styles["Code"],    fontSize=8,  leading=12, backColor=colors.HexColor("#f3f4f6"), leftIndent=12, spaceAfter=6)
        style_toc    = ParagraphStyle("RPToc",    parent=styles["Normal"],  fontSize=11, leading=20, leftIndent=12)

        story = []

        # Cover
        story.append(Spacer(1, 40 * mm))
        story.append(Paragraph("Complete Migration Assessment", style_title))
        story.append(Paragraph("Generated by Artha Talend Migration Accelerator", style_sub))
        story.append(Paragraph(datetime.now().strftime("%Y-%m-%d %H:%M:%S"), style_sub))
        story.append(PageBreak())

        # TOC
        story.append(Paragraph("Table of Contents", style_h1))
        story.append(HRFlowable(width="100%", color=colors.HexColor("#1a3c6e")))
        story.append(Spacer(1, 6))
        for idx, title in enumerate(sections.keys(), start=1):
            story.append(Paragraph(f"{idx}.  {title}", style_toc))
        story.append(PageBreak())

        # Sections
        for title, content in sections.items():
            story.append(Paragraph(title, style_h1))
            story.append(HRFlowable(width="100%", color=colors.HexColor("#1a3c6e")))
            story.append(Spacer(1, 4))
            _add_pdf_content(story, str(content or ""), style_body, style_h2, style_h3, style_bullet, style_code)
            story.append(PageBreak())

        doc.build(story)
    except Exception as e:
        # Last-resort fallback
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas as rl_canvas
            c = rl_canvas.Canvas(pdf_path, pagesize=A4)
            c.setFont("Helvetica-Bold", 16)
            c.drawString(72, 750, "Complete Migration Assessment")
            c.setFont("Helvetica", 10)
            c.drawString(72, 720, f"PDF generation error: {e}")
            c.save()
        except Exception:
            pass
    return pdf_path


def _add_pdf_content(story, content, style_body, style_h2, style_h3, style_bullet, style_code):
    """Parse markdown-ish content and append reportlab flowables."""
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors as rl_colors
    import html as html_mod

    lines = content.splitlines()
    in_code = False
    code_lines = []
    skip_first_h1 = False
    table_buffer: List[List[str]] = []

    palette_header_bg = rl_colors.HexColor("#1a3c6e")
    palette_alt_row = rl_colors.HexColor("#f8fafc")
    palette_border = rl_colors.HexColor("#e5e7eb")

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer:
            return
        # remove separator rows
        data_rows = [
            r for r in table_buffer
            if not all(set(c.replace("-", "").replace(":", "").replace(" ", "")) == set() for c in r)
        ]
        if not data_rows:
            table_buffer = []
            return
        ncols = max(len(r) for r in data_rows)
        padded = [r + [""] * (ncols - len(r)) for r in data_rows]
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        col_w = (A4[0] - 40 * mm) / ncols
        tbl = Table(padded, colWidths=[col_w] * ncols, repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0), (-1, 0), palette_header_bg),
            ("TEXTCOLOR",     (0, 0), (-1, 0), rl_colors.white),
            ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",      (0, 0), (-1, -1), 8),
            ("GRID",          (0, 0), (-1, -1), 0.4, palette_border),
            ("ROWBACKGROUNDS",(0, 1), (-1, -1), [rl_colors.white, palette_alt_row]),
            ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING",    (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 6))
        table_buffer = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("| ") or (stripped.startswith("|") and "|" in stripped[1:]):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            table_buffer.append(cells)
            continue
        else:
            flush_table()

        if stripped.startswith("```"):
            if in_code:
                if code_lines:
                    story.append(Paragraph("<br/>".join(html_mod.escape(l) for l in code_lines), style_code))
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if stripped.startswith("# ") and not skip_first_h1:
            skip_first_h1 = True
            continue
        elif stripped.startswith("# "):
            story.append(Paragraph(html_mod.escape(stripped[2:]), style_h2))
        elif stripped.startswith("## "):
            story.append(Paragraph(html_mod.escape(stripped[3:]), style_h2))
        elif stripped.startswith("### "):
            story.append(Paragraph(html_mod.escape(stripped[4:]), style_h3))
        elif stripped.startswith("- "):
            story.append(Paragraph(f"• &nbsp;{html_mod.escape(stripped[2:])}", style_bullet))
        elif stripped == "":
            story.append(Spacer(1, 4))
        else:
            story.append(Paragraph(html_mod.escape(stripped), style_body))

    flush_table()
    if in_code and code_lines:
        story.append(Paragraph("<br/>".join(html_mod.escape(l) for l in code_lines), style_code))


def _repository_overview(all_jobs: Sequence[Dict[str, Any]], repository_path: str = None) -> str:
    inventory = InventoryParser().build_inventory(all_jobs, repository_path)
    kpis = inventory.get("kpis", {})
    rows = [
        "# Repository Overview",
        "",
        f"- Total Jobs: {kpis.get('total_jobs', len(all_jobs))}",
        f"- Total Components: {kpis.get('total_components', _component_count(all_jobs))}",
        f"- Context Groups: {kpis.get('total_contexts', 'Not available')}",
        f"- Routines: {kpis.get('total_routines', 'Not available')}",
        f"- Joblets: {kpis.get('total_joblets', 'Not available')}",
        "",
        "## Jobs",
    ]
    for job in all_jobs:
        data = job.get("job_data", job)
        rows.append(f"- {data.get('job_name', 'Unknown')}: {len(data.get('components', []))} components")
    return "\n".join(rows)


def _readiness_scores(all_jobs: Sequence[Dict[str, Any]], repository_path: str = None) -> str:
    scores = RepositoryScoring().score(all_jobs, repository_path)
    labels = {
        "migration_readiness_score": "Migration Readiness",
        "cloud_readiness_score": "Cloud Readiness",
        "repository_complexity_score": "Repository Complexity",
        "documentation_readiness_score": "Documentation Readiness",
        "testing_readiness_score": "Testing Readiness",
    }
    lines = ["# Readiness Scores", ""]
    for key, label in labels.items():
        lines.append(f"- {label}: {scores.get(key, 0)}%")
    return "\n".join(lines)


def _recommendations(
    all_jobs: Sequence[Dict[str, Any]],
    auto_fix_recs: Sequence[Dict[str, Any]] = None,
) -> str:
    lines = ["# Recommendations", ""]

    ai_items = [
        (job.get("job_data", {}).get("job_name", "Unknown"), str(job.get("ai_recommendation", "")).strip())
        for job in all_jobs
        if str(job.get("ai_recommendation", "")).strip()
    ]
    if ai_items:
        lines.append("## AI Migration Recommendations")
        for job_name, recommendation in ai_items:
            lines.extend([f"### {job_name}", recommendation, ""])

    recs = list(auto_fix_recs or [])
    if recs:
        lines.append("## Remediation Recommendations")
        for rec in recs[:50]:
            title = rec.get("title") or rec.get("component") or rec.get("job_name") or "Recommendation"
            action = rec.get("action") or rec.get("recommendation") or rec.get("description") or ""
            lines.append(f"- {title}: {action}".rstrip(": "))

    if len(lines) <= 2:
        lines.extend([
            "- Replace deprecated components with supported Talend 8 equivalents.",
            "- Validate context values, connection metadata, and runtime dependencies before import.",
            "- Prioritize high-risk custom Java, shell execution, and unknown components.",
            "- Build regression tests for key source-to-target flows before cutover.",
        ])
    return "\n".join(lines)


def _merge_repository_ai_context(sections: Dict[str, str], context: Dict[str, Any]) -> Dict[str, str]:
    if not context:
        return sections
    mapping = {
        "Executive Summary": "executive_summary",
        "Repository Overview": "repository_overview",
        "Upgrade Path": "upgrade_path",
        "Readiness Scores": "readiness_scores",
        "Technical Flowchart": "technical_flowchart_notes",
        "Business Flowchart": "business_flowchart_notes",
        "Repository Flowchart": "repository_flowchart_notes",
        "Technical Documentation": "technical_documentation_notes",
        "Functional Documentation": "functional_documentation_notes",
        "KT Documentation": "kt_documentation_notes",
        "Migration Assessment": "migration_assessment",
        "Recommendations": "recommendations",
        "Routine Assessment": "routine_assessment",
        "Joblet Assessment": "joblet_assessment",
        "Java Risk Assessment": "java_risk",
        "Test Cases": "test_cases",
        "Documentation Readiness": "doc_readiness",
    }
    merged = dict(sections)
    for section, context_key in mapping.items():
        value = str(context.get(context_key, "")).strip()
        if value:
            merged[section] = f"{merged.get(section, '')}\n\n## Repository AI Context\n{value}"
    return merged


def _component_count(all_jobs: Sequence[Dict[str, Any]]) -> int:
    return sum(len(job.get("job_data", job).get("components", [])) for job in all_jobs)


def _add_markdown_content(document: Document, content: Any) -> None:
    lines = str(content or "").splitlines() or ["No content generated."]
    in_code = False
    code_lines: List[str] = []
    skipped_initial_heading = False

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                _add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if (
            not skipped_initial_heading
            and stripped.startswith("# ")
        ):
            skipped_initial_heading = True
            continue

        if stripped:
            skipped_initial_heading = True

        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif stripped:
            document.add_paragraph(stripped)
        else:
            document.add_paragraph("")

    if in_code:
        _add_code_block(document, code_lines)


def _add_code_block(document: Document, lines: Iterable[str]) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(9)
