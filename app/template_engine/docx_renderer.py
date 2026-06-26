from docx import Document

class DocxRenderer:

    @staticmethod
    def _replace_paragraph(paragraph, placeholder, value):
        full_text = "".join(r.text for r in paragraph.runs)
        if placeholder not in full_text:
            return
        new_text = full_text.replace(placeholder, str(value))
        for i, run in enumerate(paragraph.runs):
            run.text = new_text if i == 0 else ""

    def _process_paragraphs(self, paragraphs, content_map):
        for p in paragraphs:
            for ph, val in content_map.items():
                self._replace_paragraph(p, ph, val)

    def render(self, template_path, content_map, output_path):
        doc = Document(template_path)
        self._process_paragraphs(doc.paragraphs, content_map)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self._process_paragraphs(cell.paragraphs, content_map)
        doc.save(output_path)
        return output_path
